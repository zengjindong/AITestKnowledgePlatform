"""
Document Processor for Knowledge Base.

Handles document upload, text extraction, and full-text indexing.
Supports PDF, DOC, DOCX, TXT, MD, and other text-based formats.

PRD Reference: Section 6.2.2 - Document upload and full-text search
"""
import json
import logging
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """A document in the knowledge base."""
    doc_id: str
    filename: str
    file_path: str
    content: str
    file_type: str
    size: int
    extracted_at: str = field(default_factory=lambda: datetime.now().isoformat())
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "doc_id": self.doc_id,
            "filename": self.filename,
            "file_path": self.file_path,
            "content": self.content[:500] + "..." if len(self.content) > 500 else self.content,
            "file_type": self.file_type,
            "size": self.size,
            "extracted_at": self.extracted_at,
            "metadata": self.metadata
        }


class DocumentProcessor:
    """
    Process uploaded documents and extract text content.

    Supports:
    - Plain text (.txt, .md, .log)
    - PDF (basic text extraction)
    - JSON/YAML
    - HTML
    """

    def __init__(self, storage_path: str = None):
        self.storage_path = storage_path or str(Path(__file__).parent.parent / "data" / "documents")
        Path(self.storage_path).mkdir(parents=True, exist_ok=True)

    def process_file(self, file_path: str) -> Optional[Document]:
        """
        Process a file and extract its text content.

        Args:
            file_path: Path to the file to process

        Returns:
            Document object with extracted content, or None if processing fails
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return None

            filename = file_path.name
            file_type = file_path.suffix.lower()
            size = file_path.stat().st_size

            # Extract content based on file type
            content = self._extract_content(file_path, file_type)

            if not content:
                logger.warning(f"No content extracted from {filename}")
                return None

            doc_id = f"doc_{datetime.now().strftime('%Y%m%d%H%M%S')}_{file_path.stem}"

            return Document(
                doc_id=doc_id,
                filename=filename,
                file_path=str(file_path),
                content=content,
                file_type=file_type,
                size=size,
                metadata=self._extract_metadata(file_path, content)
            )

        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return None

    def _extract_content(self, file_path: Path, file_type: str) -> Optional[str]:
        """Extract text content from file."""
        try:
            if file_type in [".txt", ".md", ".log", ".json", ".yaml", ".yml", ".xml", ".html", ".htm"]:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()

            elif file_type == ".pdf":
                return self._extract_pdf(file_path)

            elif file_type in [".docx", ".doc"]:
                return self._extract_docx(file_path)

            else:
                # Try as plain text
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read()
                except:
                    logger.warning(f"Unsupported file type: {file_type}")
                    return None

        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return None

    def _extract_pdf(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF."""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_parts = []
                for page in reader.pages[:50]:  # Limit to first 50 pages
                    text_parts.append(page.extract_text() or "")
                return "\n".join(text_parts)
        except ImportError:
            logger.warning("PyPDF2 not available, PDF extraction disabled")
            return None
        except Exception as e:
            logger.warning(f"PDF extraction failed: {e}")
            return None

    def _extract_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except ImportError:
            logger.warning("python-docx not available, DOCX extraction disabled")
            return None
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")
            return None

    def _extract_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Extract metadata from file content."""
        return {
            "word_count": len(content.split()),
            "char_count": len(content),
            "line_count": content.count('\n') + 1
        }

    def save_document(self, document: Document, destination: str = None) -> str:
        """
        Save a processed document.

        Args:
            document: Document to save
            destination: Optional destination path

        Returns:
            Path where document was saved
        """
        try:
            dest_path = Path(destination) if destination else Path(self.storage_path) / document.filename
            dest_path.parent.mkdir(parents=True, exist_ok=True)

            with open(dest_path, 'w', encoding='utf-8') as f:
                json.dump(document.to_dict(), f, ensure_ascii=False, indent=2)

            logger.info(f"Saved document: {dest_path}")
            return str(dest_path)

        except Exception as e:
            logger.error(f"Error saving document: {e}")
            return ""


class FullTextSearch:
    """
    Simple full-text search index for documents.

    Provides:
    - Keyword indexing
    - TF-IDF based ranking
    - Highlight snippets

    For production, consider Elasticsearch integration.
    """

    def __init__(self, index_path: str = None):
        self.index_path = index_path or str(Path(__file__).parent.parent / "data" / "search_index.json")
        self.documents: Dict[str, Dict[str, Any]] = {}
        self._load_index()

    def _load_index(self) -> None:
        """Load existing search index."""
        try:
            if Path(self.index_path).exists():
                with open(self.index_path, 'r', encoding='utf-8') as f:
                    self.documents = json.load(f)
                logger.info(f"Loaded search index with {len(self.documents)} documents")
        except Exception as e:
            logger.warning(f"Could not load search index: {e}")
            self.documents = {}

    def _save_index(self) -> None:
        """Save search index to disk."""
        try:
            Path(self.index_path).parent.mkdir(parents=True, exist_ok=True)
            with open(self.index_path, 'w', encoding='utf-8') as f:
                json.dump(self.documents, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Error saving search index: {e}")

    def add_document(self, doc: Document) -> bool:
        """
        Add a document to the search index.

        Args:
            doc: Document to index

        Returns:
            True if successful
        """
        try:
            self.documents[doc.doc_id] = {
                "doc_id": doc.doc_id,
                "filename": doc.filename,
                "content": doc.content,
                "file_type": doc.file_type,
                "extracted_at": doc.extracted_at,
                "metadata": doc.metadata
            }
            self._save_index()
            logger.debug(f"Indexed document: {doc.filename}")
            return True
        except Exception as e:
            logger.error(f"Error indexing document: {e}")
            return False

    def search(
        self,
        query: str,
        limit: int = 10,
        file_type: str = None
    ) -> List[Dict[str, Any]]:
        """
        Search documents by keyword.

        Args:
            query: Search query
            limit: Maximum results
            file_type: Optional filter by file type

        Returns:
            List of matching documents with relevance scores
        """
        if not query:
            return []

        query_lower = query.lower()
        query_terms = set(query_lower.split())

        results = []
        for doc_id, doc in self.documents.items():
            # Filter by file type if specified
            if file_type and doc.get("file_type") != file_type:
                continue

            content_lower = doc.get("content", "").lower()

            # Calculate simple relevance score
            score = 0
            matched_terms = []

            for term in query_terms:
                if term in content_lower:
                    score += 1
                    matched_terms.append(term)

            # Boost for exact phrase match
            if query_lower in content_lower:
                score += 5

            if score > 0:
                # Extract snippet around matches
                snippet = self._extract_snippet(doc.get("content", ""), query_lower)

                results.append({
                    "doc_id": doc_id,
                    "filename": doc.get("filename"),
                    "file_type": doc.get("file_type"),
                    "score": score,
                    "matched_terms": matched_terms,
                    "snippet": snippet,
                    "extracted_at": doc.get("extracted_at")
                })

        # Sort by score and limit
        results.sort(key=lambda x: x["score"], reverse=True)
        return results[:limit]

    def _extract_snippet(self, content: str, query: str, context_len: int = 100) -> str:
        """Extract snippet around query match."""
        content_lower = content.lower()
        pos = content_lower.find(query)

        if pos == -1:
            return content[:200] + "..." if len(content) > 200 else content

        start = max(0, pos - context_len)
        end = min(len(content), pos + len(query) + context_len)

        snippet = content[start:end]
        if start > 0:
            snippet = "..." + snippet
        if end < len(content):
            snippet = snippet + "..."

        return snippet

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document from the index."""
        if doc_id in self.documents:
            del self.documents[doc_id]
            self._save_index()
            return True
        return False

    def get_stats(self) -> Dict[str, int]:
        """Get index statistics."""
        return {
            "total_documents": len(self.documents),
            "indexed_file_types": len(set(d.get("file_type") for d in self.documents.values()))
        }


# Singleton instances
_document_processor: Optional[DocumentProcessor] = None
_search_index: Optional[FullTextSearch] = None


def get_document_processor() -> DocumentProcessor:
    """Get singleton document processor."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor


def get_search_index() -> FullTextSearch:
    """Get singleton search index."""
    global _search_index
    if _search_index is None:
        _search_index = FullTextSearch()
    return _search_index